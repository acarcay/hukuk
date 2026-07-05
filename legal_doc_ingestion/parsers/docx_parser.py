"""
DOCX parser using python-docx.

Treats each paragraph as part of a single logical page unless
section/page-break information is found.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import List, Set

from legal_doc_ingestion.exceptions import CorruptedFileError, ParsingError
from legal_doc_ingestion.parsers.base import BaseParser, RawDocument, RawPage

logger = logging.getLogger(__name__)

# XML namespace used for Word page-break elements
_OOXML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class DOCXParser(BaseParser):
    """
    Extracts text from ``.docx`` files using ``python-docx``.

    The parser attempts to detect page breaks (``w:br`` with ``w:type="page"``)
    to split the document into logical pages.  When no page breaks are found,
    the entire document is returned as a single page.
    """

    @property
    def supported_extensions(self) -> Set[str]:
        return {".docx"}

    def parse(self, filepath: Path) -> RawDocument:
        try:
            from docx import Document as DocxDocument
            from docx.opc.exceptions import PackageNotFoundError
            from docx.oxml.ns import qn
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install it with: pip install python-docx"
            ) from exc

        filepath = filepath.resolve()

        try:
            doc = DocxDocument(str(filepath))
        except PackageNotFoundError as exc:
            raise CorruptedFileError(str(filepath), reason="Not a valid DOCX package.") from exc
        except Exception as exc:
            raise CorruptedFileError(str(filepath), reason=str(exc)) from exc

        raw_doc = RawDocument()

        # Attempt to read core properties
        try:
            props = doc.core_properties
            raw_doc.author = props.author or None
            raw_doc.title = props.title or None
            if props.created:
                raw_doc.creation_date = props.created.isoformat()
        except Exception:
            raw_doc.warnings.append("Could not read document core properties.")

        # ------------------------------------------------------------------
        # Build a lookup of table elements by their XML identity so we can
        # process paragraphs AND tables in document order (not separately).
        # doc.paragraphs only returns <w:p> elements; doc.tables returns
        # <w:tbl> elements — iterating doc.element.body covers both.
        # ------------------------------------------------------------------
        from docx.table import Table as DocxTable
        from docx.text.paragraph import Paragraph as DocxParagraph

        pages_text: List[str] = []
        current_page_lines: List[str] = []

        try:
            for child in doc.element.body:
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

                if tag == "p":  # paragraph
                    para = DocxParagraph(child, doc)
                    has_page_break = self._paragraph_has_page_break(para)
                    if has_page_break and current_page_lines:
                        pages_text.append("\n".join(current_page_lines))
                        current_page_lines = []
                    line = para.text
                    if line.strip():
                        current_page_lines.append(line)

                elif tag == "tbl":  # table — critical for contracts
                    table = DocxTable(child, doc)
                    table_lines: List[str] = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        # Deduplicate merged cells (python-docx repeats them)
                        deduped: List[str] = []
                        for cell in cells:
                            if not deduped or cell != deduped[-1]:
                                deduped.append(cell)
                        row_text = "\t".join(deduped)
                        if row_text.strip():
                            table_lines.append(row_text)
                    if table_lines:
                        # Wrap table content with markers so chunker can
                        # keep it together and metadata can reflect it
                        current_page_lines.append("[TABLO]")
                        current_page_lines.extend(table_lines)
                        current_page_lines.append("[/TABLO]")

            # Flush remaining content
            if current_page_lines:
                pages_text.append("\n".join(current_page_lines))
        except Exception as exc:
            raise ParsingError(str(filepath), reason=str(exc)) from exc

        if not pages_text:
            raise ParsingError(str(filepath), reason="No text content found in document.")

        for idx, text in enumerate(pages_text, start=1):
            raw_doc.pages.append(
                RawPage(page_number=idx, text=text)
            )

        return raw_doc

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _paragraph_has_page_break(para) -> bool:  # type: ignore[no-untyped-def]
        """Check if a python-docx Paragraph contains a page-break element."""
        try:
            from lxml import etree  # python-docx bundles lxml

            for run in para.runs:
                br_elements = run._element.findall(f"{{{_OOXML_NS}}}br")
                for br in br_elements:
                    if br.get(f"{{{_OOXML_NS}}}type") == "page":
                        return True
        except Exception:
            pass
        return False
