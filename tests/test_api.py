"""
Tests for the Legal RAG API endpoints.

Uses FastAPI TestClient with mocked services to avoid
needing a running Ollama instance or real embedding model.
"""

from __future__ import annotations

import io
import json
from typing import AsyncGenerator, Dict, List, Optional
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from fastapi.testclient import TestClient


# ------------------------------------------------------------------
# Mock services before importing the app
# ------------------------------------------------------------------

class MockEmbedder:
    def embed(self, text: str) -> List[float]:
        h = hash(text)
        return [(h >> i & 0xFF) / 255.0 for i in range(8)]

    def embed_batch(self, texts) -> List[List[float]]:
        return [self.embed(t) for t in texts]


class MockStore:
    def __init__(self):
        self._count = 0
        self._docs: List[Dict] = []

    @property
    def count(self) -> int:
        return self._count

    def insert_chunks(self, chunks, engine, **kwargs) -> int:
        self._count += len(chunks)
        for c in chunks:
            self._docs.append({
                "text": c.text,
                "source_id": c.source_id,
                "section_heading": c.section_heading,
                "disk_filename": c.metadata.get("disk_filename", c.source_id),
            })
        return len(chunks)

    def delete_by_source(self, source_id: str) -> None:
        removed = [d for d in self._docs if d["source_id"] == source_id]
        self._docs = [d for d in self._docs if d["source_id"] != source_id]
        self._count -= len(removed)
        if self._count < 0:
            self._count = 0

    def search(self, query, engine, k=5, **kwargs):
        from legal_doc_ingestion.vectorization.store import SearchResult
        source_filter = kwargs.get("source_filter")
        docs = self._docs
        if source_filter:
            if isinstance(source_filter, list):
                docs = [d for d in docs if d["source_id"] in source_filter]
            else:
                docs = [d for d in docs if d["source_id"] == source_filter]
        results = []
        for doc in docs[:k]:
            results.append(SearchResult(
                chunk_id="mock_id",
                text=doc["text"],
                distance=0.15,
                source_id=doc.get("source_id", "test.pdf"),
                section_heading=doc.get("section_heading"),
                metadata={
                    "source_id": doc.get("source_id", "test.pdf"),
                    "section_heading": doc.get("section_heading"),
                    "disk_filename": doc.get("disk_filename", ""),
                },
            ))
        return results

    def get_all(self, *, where=None, limit=None, include=None):
        docs = self._docs
        if where and "source_id" in where and isinstance(where["source_id"], str):
            docs = [d for d in docs if d["source_id"] == where["source_id"]]
        return {
            "ids": [f"id_{i}" for i in range(len(docs))],
            "documents": [d["text"] for d in docs],
            "metadatas": [
                {
                    "source_id": d["source_id"],
                    "document_type": "pdf",
                    "section_heading": d.get("section_heading"),
                    "disk_filename": d.get("disk_filename", d["source_id"]),
                }
                for d in docs
            ],
        }

    def get_source_metadata(self, source_id: str):
        raw = self.get_all(where={"source_id": source_id}, limit=1, include=["metadatas"])
        metas = raw.get("metadatas") or []
        return metas[0] if metas else None


class MockLLM:
    @property
    def model(self) -> str:
        return "mock-llama3:8b"

    async def aclose(self) -> None:
        pass

    async def health_check(self) -> bool:
        return True

    async def generate_stream(self, prompt, system="", **kwargs) -> AsyncGenerator:
        tokens = ["Bu ", "sorunun ", "cevabı ", "şudur."]
        for t in tokens:
            yield t

    async def generate(self, prompt, system="", **kwargs) -> str:
        return "Bu sorunun cevabı bağlamda şu şekilde açıklanmaktadır."


# ------------------------------------------------------------------
# Fixtures
# ------------------------------------------------------------------

@pytest.fixture
def mock_services():
    """Patch the global services container with mocks.

    We also patch ``services.initialize`` so that the FastAPI lifespan
    event (which calls it on TestClient startup) does NOT overwrite the
    mocks we set here with real ChromaDB / Embedder / LLM instances.
    """
    from unittest.mock import patch as _patch

    from api.dependencies import services
    from legal_doc_ingestion.ingestion import LegalDocumentIngestor
    from legal_doc_ingestion.vectorization.chunker import LegalSemanticChunker

    # Install mocks before the lifespan runs
    services.ingestor = LegalDocumentIngestor()
    services.chunker = LegalSemanticChunker(overlap_ratio=0.10, min_chunk_chars=50)
    services.embedder = MockEmbedder()
    services.store = MockStore()
    services.llm = MockLLM()

    # Patch initialize() to a no-op so TestClient lifespan cannot overwrite
    with _patch.object(services, "initialize", return_value=None):
        yield services

    # Restore to uninitialized state after test
    services.ingestor = None
    services.chunker = None
    services.embedder = None
    services.store = None
    services.llm = None


@pytest.fixture
def client(mock_services):
    """Create a TestClient with mocked services."""
    from api.main import app
    with TestClient(app, raise_server_exceptions=False) as c:
        yield c


# ------------------------------------------------------------------
# Health endpoint
# ------------------------------------------------------------------

class TestHealthEndpoint:

    def test_health_returns_200(self, client):
        resp = client.get("/health")
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "healthy"
        assert "vector_store" in data
        assert "llm" in data

    def test_root_returns_info(self, client):
        resp = client.get("/")
        assert resp.status_code == 200
        assert "Legal RAG API" in resp.json()["service"]


# ------------------------------------------------------------------
# Upload endpoint
# ------------------------------------------------------------------

class TestUploadEndpoint:

    def test_upload_rejects_unsupported_format(self, client):
        file_content = b"some content"
        resp = client.post(
            "/api/v1/upload",
            files=[("files", ("test.txt", io.BytesIO(file_content), "text/plain"))],
        )
        assert resp.status_code == 400
        assert "Unsupported" in resp.json()["detail"]

    def test_upload_rejects_empty_request(self, client):
        resp = client.post("/api/v1/upload")
        assert resp.status_code == 422  # validation error

    def test_upload_docx_success(self, client, mock_services, tmp_path):
        """Upload a real DOCX file and verify the pipeline runs."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        # Create a test DOCX
        doc = DocxDocument()
        doc.add_paragraph("MADDE 1 - Taraflar")
        doc.add_paragraph("Bu sözleşme aşağıdaki taraflar arasında imzalanmıştır.")
        doc.add_paragraph("MADDE 2 - Konu")
        doc.add_paragraph("Sözleşmenin konusu burada belirtilmiştir.")

        docx_path = tmp_path / "test_contract.docx"
        doc.save(str(docx_path))
        content = docx_path.read_bytes()

        resp = client.post(
            "/api/v1/upload",
            files=[("files", ("test_contract.docx", io.BytesIO(content),
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["total_chunks"] > 0
        assert data["documents"][0]["status"] == "success"
        assert mock_services.store.count > 0


# ------------------------------------------------------------------
# Chat endpoint
# ------------------------------------------------------------------

class TestChatEndpoint:

    def _seed_store(self, mock_services):
        """Insert some chunks so the store isn't empty."""
        from legal_doc_ingestion.vectorization.chunker import TextChunk
        chunks = [
            TextChunk(
                chunk_id="c1", text="Aylık kira bedeli 15.000 TL olarak belirlenmiştir.",
                source_id="kira.pdf", section_heading="MADDE 3",
            ),
            TextChunk(
                chunk_id="c2", text="Sözleşme süresi 1 yıl olup başlangıç tarihi 01.01.2026'dır.",
                source_id="kira.pdf", section_heading="MADDE 4",
            ),
        ]
        mock_services.store.insert_chunks(chunks, mock_services.embedder)

    def test_chat_no_documents_returns_404(self, client):
        resp = client.post(
            "/api/v1/chat",
            json={"query": "Kira bedeli nedir?", "stream": False},
        )
        assert resp.status_code == 404

    def test_chat_non_streaming(self, client, mock_services):
        self._seed_store(mock_services)
        resp = client.post(
            "/api/v1/chat",
            json={"query": "Kira bedeli nedir?", "stream": False},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert "answer" in data
        assert len(data["answer"]) > 0
        assert len(data["context"]) > 0
        assert data["model"] == "mock-llama3:8b"

    def test_chat_streaming_sse(self, client, mock_services):
        self._seed_store(mock_services)
        resp = client.post(
            "/api/v1/chat",
            json={"query": "Sözleşme süresi nedir?", "stream": True},
        )
        assert resp.status_code == 200
        assert "text/event-stream" in resp.headers.get("content-type", "")

        content = resp.text
        assert "event: context" in content
        assert "event: token" in content
        assert "event: done" in content

    def test_chat_empty_query_rejected(self, client):
        resp = client.post(
            "/api/v1/chat",
            json={"query": "", "stream": False},
        )
        assert resp.status_code == 422

    def test_chat_with_source_filter(self, client, mock_services):
        self._seed_store(mock_services)
        resp = client.post(
            "/api/v1/chat",
            json={
                "query": "Kira bedeli?",
                "stream": False,
                "source_filter": "kira.pdf",
            },
        )
        assert resp.status_code == 200

    def test_chat_with_temperature(self, client, mock_services):
        self._seed_store(mock_services)
        resp = client.post(
            "/api/v1/chat",
            json={
                "query": "Taraflar kimdir?",
                "stream": False,
                "temperature": 0.5,
            },
        )
        assert resp.status_code == 200


# ------------------------------------------------------------------
# Documents endpoint
# ------------------------------------------------------------------

class TestDocumentsEndpoint:

    def test_documents_empty(self, client):
        resp = client.get("/api/v1/documents")
        assert resp.status_code == 200
        assert resp.json()["total_vectors"] == 0

    def test_documents_after_insert(self, client, mock_services):
        from legal_doc_ingestion.vectorization.chunker import TextChunk
        chunks = [
            TextChunk(chunk_id="d1", text="Test", source_id="a.pdf"),
            TextChunk(chunk_id="d2", text="Test", source_id="b.pdf"),
        ]
        mock_services.store.insert_chunks(chunks, mock_services.embedder)
        resp = client.get("/api/v1/documents")
        assert resp.status_code == 200
        assert resp.json()["total_vectors"] == 2

    def test_delete_document(self, client, mock_services):
        """DELETE /api/v1/documents/{source_id} vektörleri siler."""
        from legal_doc_ingestion.vectorization.chunker import TextChunk
        chunks = [
            TextChunk(chunk_id="x1", text="Kira sözleşmesi", source_id="kira.pdf"),
            TextChunk(chunk_id="x2", text="Ödeme şartları", source_id="kira.pdf"),
        ]
        mock_services.store.insert_chunks(chunks, mock_services.embedder)
        assert mock_services.store.count == 2

        resp = client.delete("/api/v1/documents/kira.pdf")
        assert resp.status_code == 200
        assert resp.json()["status"] == "success"
        assert mock_services.store.count == 0

    def test_delete_nonexistent_document_still_200(self, client):
        """Olmayan bir source_id için DELETE yine de 200 dönmeli (idempotent)."""
        resp = client.delete("/api/v1/documents/hayalet.pdf")
        assert resp.status_code == 200


# ------------------------------------------------------------------
# Upload güvenlik & yeniden yükleme testleri
# ------------------------------------------------------------------

class TestUploadSecurity:

    def test_upload_rejects_oversized_file(self, client, mock_services):
        """MAX_UPLOAD_SIZE_MB sınırını aşan dosya 413 döndürmeli."""
        from api.config import settings
        import tempfile

        # Limit'i geçici olarak 0 MB yap
        original = settings.MAX_UPLOAD_SIZE_MB
        settings.MAX_UPLOAD_SIZE_MB = 0
        try:
            content = b"PDF content that is definitely > 0 MB"
            resp = client.post(
                "/api/v1/upload",
                files=[("files", ("test.pdf", io.BytesIO(content), "application/pdf"))],
            )
            assert resp.status_code == 413
        finally:
            settings.MAX_UPLOAD_SIZE_MB = original

    def test_upload_rejects_txt_extension(self, client):
        """Desteklenmeyen format 400 döndürmeli."""
        resp = client.post(
            "/api/v1/upload",
            files=[("files", ("notes.txt", io.BytesIO(b"some text"), "text/plain"))],
        )
        assert resp.status_code == 400
        assert "Unsupported" in resp.json()["detail"]

    def test_upload_rejects_empty_filename(self, client):
        """Dosya listesi boş gönderilince 422 döndürmeli."""
        resp = client.post("/api/v1/upload")
        assert resp.status_code == 422


class TestReUploadBehavior:
    """Aynı kaynak belgenin yeniden yüklenmesi eski chunk'ları temizlemeli."""

    def test_reupload_clears_stale_chunks(self, client, mock_services, tmp_path):
        """
        Bug 4 regresyon testi:
        İlk yükleme → 3 chunk; ikinci yükleme (farklı içerik) → 2 chunk.
        İkinci yüklemeden sonra store'da toplam 2 chunk olmalı (3 değil).
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        def make_docx(paragraphs):
            doc = DocxDocument()
            for p in paragraphs:
                doc.add_paragraph(p)
            path = tmp_path / f"contract_{len(paragraphs)}.docx"
            doc.save(str(path))
            return path.read_bytes()

        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # İlk yükleme
        content_v1 = make_docx([
            "MADDE 1 - Taraflar", "Kiracı: Ahmet",
            "MADDE 2 - Konu", "Kiralık daire",
            "MADDE 3 - Ödeme", "15.000 TL",
        ])
        resp1 = client.post(
            "/api/v1/upload",
            files=[("files", ("sozlesme.docx", io.BytesIO(content_v1), mime))],
        )
        assert resp1.status_code == 200
        chunks_v1 = resp1.json()["total_chunks"]
        assert chunks_v1 > 0

        # İkinci yükleme (aynı dosya adı, daha az madde)
        content_v2 = make_docx([
            "MADDE 1 - Taraflar", "Kiracı: Mehmet",
            "MADDE 2 - Konu", "Ofis kirası",
        ])
        resp2 = client.post(
            "/api/v1/upload",
            files=[("files", ("sozlesme.docx", io.BytesIO(content_v2), mime))],
        )
        assert resp2.status_code == 200
        chunks_v2 = resp2.json()["total_chunks"]
        assert chunks_v2 > 0

        count_after_v2 = mock_services.store.count
        # Yeniden yükleme önce delete_by_source çağırmalı,
        # bu yüzden v1 chunk'ları temizlenmeli ve sadece v2 chunk'ları kalmalı
        assert count_after_v2 == chunks_v2, (
            f"Stale chunk'lar temizlenmemiş: store'da {count_after_v2} var, "
            f"beklenen {chunks_v2}"
        )


# ------------------------------------------------------------------
# Health endpoint edge cases
# ------------------------------------------------------------------

class TestHealthEdgeCases:

    def test_health_includes_vector_count(self, client, mock_services):
        """Health endpoint vector_store bilgisi içermeli."""
        resp = client.get("/health")
        assert resp.status_code == 200
        data = resp.json()
        assert "vector_store" in data
        assert "llm" in data
        assert "status" in data

    def test_request_id_propagated_in_response(self, client):
        """X-Request-ID header'ı response'a yansımalı."""
        resp = client.get("/health", headers={"X-Request-ID": "test-123"})
        assert resp.headers.get("X-Request-ID") == "test-123"


# ------------------------------------------------------------------
# API-key authentication
# ------------------------------------------------------------------

class TestApiKeyAuth:
    """require_api_key: anahtar tanımlıysa hassas endpoint'ler korunmalı."""

    def test_protected_endpoint_401_without_key(self, client, mock_services):
        """API_KEY set edilince anahtarsız istek 401 dönmeli."""
        from api.config import settings

        original = settings.API_KEY
        settings.API_KEY = "s3cret"
        try:
            resp = client.get("/api/v1/documents")  # no X-API-Key header
            assert resp.status_code == 401
        finally:
            settings.API_KEY = original

    def test_protected_endpoint_401_with_wrong_key(self, client, mock_services):
        from api.config import settings

        original = settings.API_KEY
        settings.API_KEY = "s3cret"
        try:
            resp = client.get(
                "/api/v1/documents", headers={"X-API-Key": "wrong"}
            )
            assert resp.status_code == 401
        finally:
            settings.API_KEY = original

    def test_protected_endpoint_200_with_correct_key(self, client, mock_services):
        from api.config import settings

        original = settings.API_KEY
        settings.API_KEY = "s3cret"
        try:
            resp = client.get(
                "/api/v1/documents", headers={"X-API-Key": "s3cret"}
            )
            assert resp.status_code == 200
        finally:
            settings.API_KEY = original

    def test_auth_disabled_allows_access(self, client):
        """API_KEY boşken (varsayılan) auth devre dışı — 200."""
        resp = client.get("/api/v1/documents")
        assert resp.status_code == 200

    def test_health_not_protected(self, client, mock_services):
        """/health auth gerektirmemeli (anahtar set olsa bile)."""
        from api.config import settings

        original = settings.API_KEY
        settings.API_KEY = "s3cret"
        try:
            resp = client.get("/health")
            assert resp.status_code == 200
        finally:
            settings.API_KEY = original
