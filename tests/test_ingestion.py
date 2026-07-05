"""
Unit tests for the legal document ingestion module.

Run with: pytest tests/ -v
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import pytest

from legal_doc_ingestion.cleaning import TextCleaner
from legal_doc_ingestion.exceptions import (
    CorruptedFileError,
    IngestionError,
    UnsupportedFormatError,
)
from legal_doc_ingestion.ingestion import LegalDocumentIngestor
from legal_doc_ingestion.models import DocumentResult, DocumentType, PageContent


# ======================================================================
# Text Cleaner Tests
# ======================================================================


class TestTextCleaner:
    """Tests for the TextCleaner utility."""

    def setup_method(self) -> None:
        self.cleaner = TextCleaner()

    def test_removes_watermark_lines(self) -> None:
        text = "Line one\nCONFIDENTIAL\nLine two\nDRAFT\nLine three"
        result = self.cleaner.clean(text)
        assert "CONFIDENTIAL" not in result
        assert "DRAFT" not in result
        assert "Line one" in result
        assert "Line two" in result
        assert "Line three" in result

    def test_removes_page_numbers_bare(self) -> None:
        text = "Some legal text.\n\n  42  \n\nMore text follows."
        result = self.cleaner.clean(text)
        assert "42" not in result
        assert "Some legal text." in result
        assert "More text follows." in result

    def test_removes_page_of_pattern(self) -> None:
        text = "Content here.\nPage 3 of 10\nMore content."
        result = self.cleaner.clean(text)
        assert "Page 3 of 10" not in result

    def test_removes_dash_page_numbers(self) -> None:
        text = "Text above.\n- 7 -\nText below."
        result = self.cleaner.clean(text)
        assert "- 7 -" not in result

    def test_collapses_excessive_newlines(self) -> None:
        text = "Para one.\n\n\n\n\nPara two."
        result = self.cleaner.clean(text)
        # Should have at most 2 consecutive newlines
        assert "\n\n\n" not in result
        assert "Para one." in result
        assert "Para two." in result

    def test_normalizes_smart_quotes(self) -> None:
        text = "\u201cHello\u201d and \u2018world\u2019"
        result = self.cleaner.clean(text)
        assert '"Hello"' in result
        assert "'world'" in result

    def test_removes_null_bytes(self) -> None:
        text = "Clean\x00 text\x00 here"
        result = self.cleaner.clean(text)
        assert "\x00" not in result
        assert "Clean text here" in result

    def test_strips_stray_urls(self) -> None:
        text = "Content\nhttps://www.example.com/footer\nMore content"
        result = self.cleaner.clean(text)
        assert "https://www.example.com" not in result

    def test_custom_watermark(self) -> None:
        cleaner = TextCleaner(extra_watermarks=[r"TASLAK"])
        text = "Madde 1.\nTASLAK\nMadde 2."
        result = cleaner.clean(text)
        assert "TASLAK" not in result

    def test_preserves_inline_watermark_words(self) -> None:
        """Watermark removal only strips lines that *solely* match the pattern."""
        text = "This is a confidential agreement between the parties."
        result = self.cleaner.clean(text)
        # The full line does NOT match the watermark pattern (fullmatch), so it stays
        assert "confidential agreement" in result

    def test_custom_step_insertion(self) -> None:
        def upper(text: str) -> str:
            return text.upper()

        self.cleaner.add_step("uppercase", upper)
        result = self.cleaner.clean("hello world")
        assert result == "HELLO WORLD"

    def test_turkish_page_number_pattern(self) -> None:
        text = "İçerik\nSayfa 3/10\nDevam"
        result = self.cleaner.clean(text)
        assert "Sayfa 3/10" not in result


# ======================================================================
# Model Tests
# ======================================================================


class TestModels:
    """Tests for Pydantic data models."""

    def test_document_result_serialization(self) -> None:
        from legal_doc_ingestion.models import DocumentMetadata

        metadata = DocumentMetadata(
            filename="test.pdf",
            filepath="/tmp/test.pdf",
            document_type=DocumentType.PDF,
            total_pages=1,
            file_size_bytes=1024,
        )
        page = PageContent(
            page_number=1,
            raw_text="Raw text",
            cleaned_text="Clean text",
        )
        result = DocumentResult(metadata=metadata, pages=[page])

        json_str = result.to_json()
        parsed = json.loads(json_str)

        assert parsed["metadata"]["filename"] == "test.pdf"
        assert parsed["metadata"]["document_type"] == "pdf"
        assert len(parsed["pages"]) == 1
        assert parsed["pages"][0]["cleaned_text"] == "Clean text"

    def test_page_count_mismatch_raises(self) -> None:
        from legal_doc_ingestion.models import DocumentMetadata

        metadata = DocumentMetadata(
            filename="test.pdf",
            filepath="/tmp/test.pdf",
            document_type=DocumentType.PDF,
            total_pages=5,  # mismatch!
            file_size_bytes=1024,
        )
        page = PageContent(
            page_number=1,
            raw_text="text",
            cleaned_text="text",
        )
        with pytest.raises(ValueError, match="Page count mismatch"):
            DocumentResult(metadata=metadata, pages=[page])

    def test_full_cleaned_text_concatenation(self) -> None:
        from legal_doc_ingestion.models import DocumentMetadata

        metadata = DocumentMetadata(
            filename="test.pdf",
            filepath="/tmp/test.pdf",
            document_type=DocumentType.PDF,
            total_pages=2,
            file_size_bytes=100,
        )
        pages = [
            PageContent(page_number=1, raw_text="A", cleaned_text="Page A"),
            PageContent(page_number=2, raw_text="B", cleaned_text="Page B"),
        ]
        result = DocumentResult(metadata=metadata, pages=pages)
        assert result.full_cleaned_text() == "Page A\n\nPage B"
        assert result.full_cleaned_text(separator=" | ") == "Page A | Page B"


# ======================================================================
# Ingestor Validation Tests
# ======================================================================


class TestIngestorValidation:
    """Tests for the LegalDocumentIngestor input validation."""

    def setup_method(self) -> None:
        self.ingestor = LegalDocumentIngestor()

    def test_nonexistent_file_raises_corrupted_error(self) -> None:
        with pytest.raises(CorruptedFileError, match="does not exist"):
            self.ingestor.ingest("/nonexistent/file.pdf")

    def test_empty_file_raises_corrupted_error(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(b"")  # 0 bytes
            path = f.name

        with pytest.raises(CorruptedFileError, match="empty"):
            self.ingestor.ingest(path)

        Path(path).unlink(missing_ok=True)

    def test_unsupported_extension_raises(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
            f.write(b"some content")
            path = f.name

        with pytest.raises(UnsupportedFormatError, match=".xyz"):
            self.ingestor.ingest(path)

        Path(path).unlink(missing_ok=True)

    def test_batch_skip_errors(self) -> None:
        """Batch ingestion with skip_errors=True should not raise."""
        results = self.ingestor.ingest_batch(
            ["/nonexistent/a.pdf", "/nonexistent/b.docx"],
            skip_errors=True,
        )
        assert len(results) == 2
        assert all(isinstance(r, IngestionError) for r in results)


# ======================================================================
# DOCX Parser Integration Test
# ======================================================================


class TestDocxParsing:
    """Integration test for DOCX parsing (requires python-docx)."""

    def test_parse_simple_docx(self) -> None:
        """Create a minimal DOCX in-memory and parse it."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        # Create a test DOCX
        doc = DocxDocument()
        doc.add_paragraph("MADDE 1 — Taraflar")
        doc.add_paragraph("Bu sözleşme aşağıdaki taraflar arasında imzalanmıştır.")
        doc.add_paragraph("MADDE 2 — Konu")
        doc.add_paragraph("Sözleşmenin konusu aşağıda belirtilmiştir.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            path = f.name

        try:
            ingestor = LegalDocumentIngestor()
            result = ingestor.ingest(path)

            assert result.metadata.document_type == DocumentType.DOCX
            assert result.metadata.filename.endswith(".docx")
            assert len(result.pages) >= 1
            assert "MADDE 1" in result.full_cleaned_text()
            assert "Taraflar" in result.full_cleaned_text()

            # Verify JSON serialization round-trip
            json_str = result.to_json()
            parsed = json.loads(json_str)
            assert "metadata" in parsed
            assert "pages" in parsed
        finally:
            Path(path).unlink(missing_ok=True)


# ======================================================================
# RTF Parser Integration Test
# ======================================================================


class TestRtfParsing:
    """Integration test for RTF parsing (requires striprtf)."""

    def test_parse_simple_rtf(self) -> None:
        try:
            from striprtf.striprtf import rtf_to_text  # noqa: F401
        except ImportError:
            pytest.skip("striprtf not installed")

        rtf_content = (
            r"{\rtf1\ansi\deff0"
            r"{\fonttbl{\f0 Times New Roman;}}"
            r"\pard Madde 1 - Taraflar\par"
            r"\pard Bu sozlesme asagidaki taraflar arasinda imzalanmistir.\par"
            r"}"
        )

        with tempfile.NamedTemporaryFile(
            suffix=".rtf", delete=False, mode="w", encoding="utf-8"
        ) as f:
            f.write(rtf_content)
            path = f.name

        try:
            ingestor = LegalDocumentIngestor()
            result = ingestor.ingest(path)

            assert result.metadata.document_type == DocumentType.RTF
            assert len(result.pages) >= 1
            assert "Madde 1" in result.full_cleaned_text()
        finally:
            Path(path).unlink(missing_ok=True)


# ======================================================================
# Exception Hierarchy Tests
# ======================================================================


class TestExceptionHierarchy:
    """Verify that all custom exceptions inherit from IngestionError."""

    def test_all_exceptions_are_ingestion_errors(self) -> None:
        from legal_doc_ingestion.exceptions import (
            CorruptedFileError,
            OCRError,
            ParsingError,
            UnsupportedFormatError,
        )

        for exc_cls in (CorruptedFileError, OCRError, ParsingError, UnsupportedFormatError):
            assert issubclass(exc_cls, IngestionError)

    def test_corrupted_file_error_attributes(self) -> None:
        exc = CorruptedFileError("/tmp/bad.pdf", reason="magic bytes invalid")
        assert exc.filepath == "/tmp/bad.pdf"
        assert "magic bytes invalid" in str(exc)


# ======================================================================
# DOCX Tablo Desteği (Bug 10 regresyon)
# ======================================================================


class TestDocxTableParsing:
    """DOCX içindeki tablolar ingestion'da görünmeli."""

    def test_table_content_included(self) -> None:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = DocxDocument()
        doc.add_paragraph("MADDE 1 - Taraflar")
        doc.add_paragraph("Bu sözleşme aşağıdaki taraflar arasında imzalanmıştır.")

        # Taraf bilgileri tablosu (sözleşmelerde çok yaygın)
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Unvan"
        table.cell(0, 1).text = "Ad Soyad"
        table.cell(1, 0).text = "Kiralayan"
        table.cell(1, 1).text = "Mehmet Yılmaz"
        table.cell(2, 0).text = "Kiracı"
        table.cell(2, 1).text = "Ayşe Demir"

        doc.add_paragraph("MADDE 2 - Ödeme")
        doc.add_paragraph("Kira bedeli aşağıdaki tabloda belirtilmiştir.")

        # Ödeme planı tablosu
        ptable = doc.add_table(rows=2, cols=3)
        ptable.cell(0, 0).text = "Ay"
        ptable.cell(0, 1).text = "Tutar"
        ptable.cell(0, 2).text = "Vade"
        ptable.cell(1, 0).text = "Ocak 2026"
        ptable.cell(1, 1).text = "15.000 TL"
        ptable.cell(1, 2).text = "05.01.2026"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            path = f.name

        try:
            ingestor = LegalDocumentIngestor()
            result = ingestor.ingest(path)
            full_text = result.full_cleaned_text()

            # Tablo içeriği metne dahil olmalı
            assert "Mehmet Yılmaz" in full_text, "Kiralayan ismi tabloda görünmeli"
            assert "Ayşe Demir" in full_text, "Kiracı ismi tabloda görünmeli"
            assert "15.000 TL" in full_text, "Kira bedeli ödeme tablosunda görünmeli"
            assert "05.01.2026" in full_text, "Vade tarihi tabloda görünmeli"
        finally:
            Path(path).unlink(missing_ok=True)

    def test_table_marker_present(self) -> None:
        """Tablo içeriği [TABLO] etiketiyle sarmalanmış olmalı."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        from legal_doc_ingestion.parsers.docx_parser import DOCXParser

        doc = DocxDocument()
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "A"
        t.cell(0, 1).text = "B"
        t.cell(1, 0).text = "C"
        t.cell(1, 1).text = "D"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            path = f.name

        try:
            parser = DOCXParser()
            raw_doc = parser.parse(Path(path))
            full_text = "\n".join(p.text for p in raw_doc.pages)
            assert "[TABLO]" in full_text
            assert "[/TABLO]" in full_text
        finally:
            Path(path).unlink(missing_ok=True)


# ======================================================================
# PDF Parser Güvenlik Limitleri (Bug 11 regresyon)
# ======================================================================


class TestPDFParserLimits:
    """PDF parser max_pages ve ocr_timeout parametrelerini doğrula."""

    def test_max_pages_constructor_parameter(self) -> None:
        """PDFParser max_pages parametresini kabul etmeli."""
        from legal_doc_ingestion.parsers.pdf_parser import PDFParser
        parser = PDFParser(max_pages=50, ocr_timeout_seconds=10)
        assert parser._max_pages == 50
        assert parser._ocr_timeout == 10

    def test_default_max_pages_is_200(self) -> None:
        from legal_doc_ingestion.parsers.pdf_parser import PDFParser
        parser = PDFParser()
        assert parser._max_pages == 200

    def test_ocr_timeout_zero_disables_timeout(self) -> None:
        """ocr_timeout_seconds=0 olduğunda _ocr_page direkt çağrılmalı (timeout bypass)."""
        from legal_doc_ingestion.parsers.pdf_parser import PDFParser
        from unittest.mock import MagicMock, patch
        from legal_doc_ingestion.parsers.base import RawDocument

        parser = PDFParser(ocr_timeout_seconds=0)
        fake_page = MagicMock()

        with patch.object(parser, "_ocr_page", return_value=("ocr text", 0.95)) as mock_ocr:
            raw_doc = RawDocument()
            result = parser._ocr_page_with_timeout(fake_page, 1, raw_doc)
            mock_ocr.assert_called_once_with(fake_page)
            assert result == ("ocr text", 0.95)
